# STUDENT ASSISTANT FEATURE
"""
Document Processing Module

Handles file upload, parsing, and text extraction for various document types.
"""

import os
import uuid
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional, List, BinaryIO, Dict, Any

from fastapi import UploadFile, HTTPException
import aiofiles

from .models import (
    Document, DocumentType, ProcessingStatus, Chapter,
    UnsupportedFileTypeError, DocumentParsingError
)


class DocumentProcessor:
    """Handles document upload and processing."""
    
    def __init__(self, upload_dir: str = "uploads/documents"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Supported file types and max sizes
        self.supported_types = {
            "application/pdf": DocumentType.PDF,
            "application/epub+zip": DocumentType.EPUB,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
            "text/plain": DocumentType.TXT,
        }
        
        self.max_file_size = 100 * 1024 * 1024  # 100MB
    
    async def upload_document(
        self, 
        file: UploadFile, 
        user_id: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Upload and process a document file.
        
        Args:
            file: Uploaded file
            user_id: ID of the user uploading
            metadata: Optional metadata (title, author, etc.)
            
        Returns:
            Document object with basic info (processing happens async)
            
        Raises:
            UnsupportedFileTypeError: If file type not supported
            HTTPException: If file too large or other validation fails
        """
        # Validate file type
        if file.content_type not in self.supported_types:
            raise UnsupportedFileTypeError(
                f"File type {file.content_type} not supported. "
                f"Supported types: {list(self.supported_types.keys())}"
            )
        
        # Validate file size
        file_size = 0
        content = await file.read()
        file_size = len(content)
        
        if file_size > self.max_file_size:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {self.max_file_size / (1024*1024):.1f}MB"
            )
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="Empty file uploaded")
        
        # Generate document ID and file path
        doc_id = str(uuid.uuid4())
        file_extension = self.supported_types[file.content_type].value
        filename = f"{doc_id}.{file_extension}"
        file_path = self.upload_dir / filename
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        # Extract basic metadata
        title = metadata.get("title") if metadata else None
        if not title:
            title = file.filename or f"Document_{doc_id[:8]}"
            # Remove file extension from title
            title = Path(title).stem
        
        # Create document object
        document = Document(
            id=doc_id,
            title=title,
            file_type=self.supported_types[file.content_type],
            file_path=str(file_path),
            file_size=file_size,
            upload_date=datetime.now(),
            author=metadata.get("author") if metadata else None,
            description=metadata.get("description") if metadata else None,
            subject=metadata.get("subject") if metadata else None,
            processing_status=ProcessingStatus.PENDING
        )
        
        return document
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = path.stat()
        
        return {
            "size": stat.st_size,
            "created": datetime.fromtimestamp(stat.st_ctime),
            "modified": datetime.fromtimestamp(stat.st_mtime),
            "extension": path.suffix.lower(),
            "name": path.name
        }
    
    def delete_document_file(self, document: Document) -> bool:
        """Delete the physical file for a document."""
        try:
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
                return True
            return False
        except Exception:
            return False
    
    def validate_document_integrity(self, document: Document) -> bool:
        """Check if document file still exists and is valid."""
        try:
            file_path = Path(document.file_path)
            if not file_path.exists():
                return False
            
            # Check file size matches
            actual_size = file_path.stat().st_size
            return actual_size == document.file_size
            
        except Exception:
            return False


# File type specific processors will be imported here
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import ebooklib
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False


class TextExtractor:
    """Extracts text content from various document types."""
    
    def __init__(self):
        self.extractors = {
            DocumentType.PDF: self._extract_pdf,
            DocumentType.DOCX: self._extract_docx,
            DocumentType.EPUB: self._extract_epub,
            DocumentType.TXT: self._extract_txt,
        }
    
    def extract_text(self, document: Document) -> Dict[str, Any]:
        """
        Extract text content from document.
        
        Returns:
            Dict with:
            - text: str - Full text content
            - chapters: List[Chapter] - Detected chapters
            - metadata: Dict - Additional metadata
            - page_count: int - Number of pages (if applicable)
        """
        if document.file_type not in self.extractors:
            raise UnsupportedFileTypeError(f"No extractor for {document.file_type}")
        
        try:
            return self.extractors[document.file_type](document.file_path)
        except Exception as e:
            raise DocumentParsingError(f"Failed to extract text: {str(e)}")
    
    def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        # Try to detect chapters by common patterns
        chapters = self._detect_chapters_in_text(text)
        
        return {
            "text": text,
            "chapters": chapters,
            "metadata": {"encoding": "utf-8"},
            "page_count": None
        }
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise DocumentParsingError("PyPDF2 not installed. Cannot process PDF files.")
        
        text_content = []
        page_count = 0
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num}]\n{page_text}\n")
                except Exception:
                    # Skip pages that can't be processed
                    continue
        
        full_text = "\n".join(text_content)
        chapters = self._detect_chapters_in_text(full_text)
        
        # Try to extract PDF metadata
        metadata = {}
        try:
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                }
        except Exception:
            pass
        
        return {
            "text": full_text,
            "chapters": chapters,
            "metadata": metadata,
            "page_count": page_count
        }
    
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise DocumentParsingError("python-docx not installed. Cannot process DOCX files.")
        
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        full_text = "\n".join(paragraphs)
        chapters = self._detect_chapters_in_text(full_text)
        
        # Extract metadata
        metadata = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "subject": doc.core_properties.subject or "",
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified,
        }
        
        return {
            "text": full_text,
            "chapters": chapters,
            "metadata": metadata,
            "page_count": None  # DOCX doesn't have fixed pages
        }
    
    def _extract_epub(self, file_path: str) -> Dict[str, Any]:
        """Extract text from EPUB file."""
        if not EPUB_AVAILABLE:
            raise DocumentParsingError("ebooklib not installed. Cannot process EPUB files.")
        
        book = epub.read_epub(file_path)
        
        chapters = []
        full_text_parts = []
        
        # Extract chapters
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Parse HTML content
                content = item.get_content().decode('utf-8')
                
                # Simple HTML tag removal (basic)
                import re
                text = re.sub(r'<[^>]+>', '', content)
                text = re.sub(r'\s+', ' ', text).strip()
                
                if text:
                    chapter_title = item.get_name() or f"Chapter {len(chapters) + 1}"
                    
                    chapter = Chapter(
                        id=str(uuid.uuid4()),
                        title=chapter_title,
                        content_preview=text[:200]
                    )
                    chapters.append(chapter)
                    full_text_parts.append(f"[{chapter_title}]\n{text}\n")
        
        full_text = "\n".join(full_text_parts)
        
        # Extract metadata
        metadata = {
            "title": book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else "",
            "author": book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else "",
            "language": book.get_metadata('DC', 'language')[0][0] if book.get_metadata('DC', 'language') else "",
        }
        
        return {
            "text": full_text,
            "chapters": chapters,
            "metadata": metadata,
            "page_count": None
        }
    
    def _detect_chapters_in_text(self, text: str) -> List[Chapter]:
        """Detect chapter boundaries in text using common patterns."""
        import re
        
        chapters = []
        
        # Common chapter patterns
        patterns = [
            r'^Chapter\s+(\d+)[:\s]*(.*)$',
            r'^CHAPTER\s+(\d+)[:\s]*(.*)$',
            r'^\d+\.\s+(.+)$',
            r'^([A-Z][A-Z\s]{10,})$',  # ALL CAPS headings
        ]
        
        lines = text.split('\n')
        current_chapter = None
        chapter_text = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check if line matches chapter pattern
            is_chapter = False
            chapter_title = None
            
            for pattern in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    if len(match.groups()) >= 2:
                        chapter_title = f"Chapter {match.group(1)}: {match.group(2)}"
                    else:
                        chapter_title = match.group(1) if match.groups() else line
                    is_chapter = True
                    break
            
            if is_chapter and chapter_title:
                # Save previous chapter
                if current_chapter and chapter_text:
                    current_chapter.content_preview = ' '.join(chapter_text)[:200]
                    chapters.append(current_chapter)
                
                # Start new chapter
                current_chapter = Chapter(
                    id=str(uuid.uuid4()),
                    title=chapter_title.strip()
                )
                chapter_text = []
            else:
                # Add to current chapter content
                if current_chapter:
                    chapter_text.append(line)
        
        # Add final chapter
        if current_chapter and chapter_text:
            current_chapter.content_preview = ' '.join(chapter_text)[:200]
            chapters.append(current_chapter)
        
        return chapters


# Availability check
